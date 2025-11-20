"""
File extraction utilities for PDF, DOCX, DOC, and TXT files
"""
import os
import io
import logging
from typing import Optional, Tuple
from fastapi import UploadFile

# PDF libraries
import pypdf
from pdfminer.high_level import extract_text as pdfminer_extract_text

# DOCX libraries
import docx

from config import settings

logger = logging.getLogger(__name__)


def sanitize_text(text: str, max_length: Optional[int] = None) -> Tuple[str, bool]:
    """
    Sanitize extracted text: remove null bytes, collapse whitespace, trim length
    
    Args:
        text: Raw extracted text
        max_length: Maximum character length (default from settings)
    
    Returns:
        Tuple of (sanitized_text, was_truncated)
    """
    if max_length is None:
        max_length = settings.MAX_TEXT_LENGTH
    
    # Remove null bytes and other problematic characters
    text = text.replace('\x00', '')
    
    # Collapse multiple whitespaces while preserving single newlines
    lines = text.split('\n')
    lines = [' '.join(line.split()) for line in lines]
    text = '\n'.join(lines)
    
    # Trim excessive newlines
    while '\n\n\n' in text:
        text = text.replace('\n\n\n', '\n\n')
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    # Truncate if necessary
    was_truncated = False
    if len(text) > max_length:
        text = text[:max_length]
        was_truncated = True
        logger.warning(f"Text truncated from original length to {max_length} chars")
    
    return text, was_truncated


def extract_text_from_pdf(file_content: bytes, filename: str) -> str:
    """
    Extract text from PDF using pypdf (primary) or pdfminer.six (fallback)
    
    Args:
        file_content: Binary content of the PDF file
        filename: Name of the file (for logging)
    
    Returns:
        Extracted text
    
    Raises:
        Exception: If extraction fails with both libraries
    """
    # Try pypdf first (simpler, faster)
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract page {page_num} from {filename}: {e}")
        
        text = '\n'.join(text_parts)
        
        # If we got reasonable text, return it
        if text and len(text.strip()) > 50:
            logger.info(f"Successfully extracted {len(text)} chars from {filename} using pypdf")
            return text
        else:
            logger.warning(f"pypdf extracted insufficient text from {filename}, trying pdfminer")
    
    except Exception as e:
        logger.warning(f"pypdf failed for {filename}: {e}, trying pdfminer")
    
    # Fallback to pdfminer.six
    try:
        text = pdfminer_extract_text(io.BytesIO(file_content))
        logger.info(f"Successfully extracted {len(text)} chars from {filename} using pdfminer")
        return text
    except Exception as e:
        logger.error(f"pdfminer also failed for {filename}: {e}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes, filename: str) -> str:
    """
    Extract text from DOCX using python-docx
    
    Args:
        file_content: Binary content of the DOCX file
        filename: Name of the file (for logging)
    
    Returns:
        Extracted text
    
    Raises:
        Exception: If extraction fails
    """
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = '\n'.join(text_parts)
        
        if text and len(text.strip()) > 20:
            logger.info(f"Successfully extracted {len(text)} chars from {filename} using python-docx")
            return text
        else:
            raise Exception("Extracted text is too short or empty")
    
    except Exception as e:
        logger.error(f"python-docx failed for {filename}: {e}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_content: bytes, filename: str) -> str:
    """
    Extract text from TXT file with encoding fallback
    
    Args:
        file_content: Binary content of the TXT file
        filename: Name of the file (for logging)
    
    Returns:
        Extracted text
    
    Raises:
        Exception: If decoding fails with both encodings
    """
    # Try UTF-8 first
    try:
        text = file_content.decode('utf-8')
        logger.info(f"Successfully decoded {filename} as UTF-8")
        return text
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
    
    # Fallback to latin-1 (never fails but may produce garbage)
    try:
        text = file_content.decode('latin-1')
        logger.info(f"Successfully decoded {filename} as latin-1")
        return text
    except Exception as e:
        logger.error(f"Failed to decode {filename}: {e}")
        raise Exception(f"Failed to decode text file: {str(e)}")


async def extract_text_from_file(upload_file: UploadFile) -> Tuple[str, Optional[str]]:
    """
    Extract text from uploaded file based on content type
    
    Args:
        upload_file: FastAPI UploadFile object
    
    Returns:
        Tuple of (extracted_text, warning_message)
    
    Raises:
        ValueError: If file type is not supported
        Exception: If extraction fails
    """
    filename = upload_file.filename or "unknown"
    content_type = upload_file.content_type or ""
    
    logger.info(f"Extracting text from {filename} (type: {content_type})")
    
    # Read file content
    file_content = await upload_file.read()
    file_size = len(file_content)
    
    # Validate file size
    if file_size > settings.MAX_FILE_SIZE:
        raise ValueError(
            f"File size ({file_size} bytes) exceeds maximum allowed size "
            f"({settings.MAX_FILE_SIZE} bytes)"
        )
    
    # Extract based on content type
    raw_text = ""
    
    if content_type == "application/pdf" or filename.lower().endswith('.pdf'):
        raw_text = extract_text_from_pdf(file_content, filename)
    
    elif content_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ] or filename.lower().endswith(('.docx', '.doc')):
        raw_text = extract_text_from_docx(file_content, filename)
    
    elif content_type == "text/plain" or filename.lower().endswith('.txt'):
        raw_text = extract_text_from_txt(file_content, filename)
    
    else:
        raise ValueError(
            f"Unsupported file type: {content_type}. "
            f"Supported types: PDF, DOCX, DOC, TXT"
        )
    
    # Sanitize and check for truncation
    clean_text, was_truncated = sanitize_text(raw_text)
    
    warning = None
    if was_truncated:
        warning = f"Text from {filename} was truncated to {settings.MAX_TEXT_LENGTH} characters"
    
    if not clean_text or len(clean_text.strip()) < 10:
        raise ValueError(f"Extracted text from {filename} is too short or empty")
    
    return clean_text, warning
